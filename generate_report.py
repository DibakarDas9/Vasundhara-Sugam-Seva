
"""
Generate Project Analysis Report (DOCX)
Reads feature_analysis_v2.md and converts it to a formatted Word document.
"""
import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx library not found.")
    print("Please install it running: pip install python-docx")
    sys.exit(1)

# Configuration
SOURCE_FILE = r'C:\Users\91896\.gemini\antigravity\brain\395e6517-04bd-433d-935c-8c2ee25b9ad0\feature_analysis_v2.md'
OUTPUT_FILE = r'c:\Users\91896\Desktop\Vasundhara- Sugam Seva\Vasundhara_Project_Report_v2.docx'

def create_report():
    if not os.path.exists(SOURCE_FILE):
        print(f"Error: Source file not found at {SOURCE_FILE}")
        return

    doc = Document()
    
    # Title
    title = doc.add_heading('Vasundhara Project Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    with open(SOURCE_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip title in MD if we already added it
        if i == 0 and line.startswith('# '):
            i += 1
            continue

        # Code Blocks
        if line.lstrip().startswith('```'):
            if in_code_block:
                # End of block
                in_code_block = False
                if code_buffer:
                    p = doc.add_paragraph('\n'.join(code_buffer))
                    p.style = 'Normal' # Use a monospaced style if available, or just Normal
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                code_buffer = []
            else:
                # Start of block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Headers
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        
        # Lists
        elif line.lstrip().startswith('* ') or line.lstrip().startswith('- '):
            text = line.lstrip()[2:]
            # Handle bold in list items
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
            
        # Tables
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i+1] and '---' in lines[i+1]:
            # Parse table
            headers = [c.strip() for c in line.strip('|').split('|')]
            i += 2 # Skip separator
            rows = []
            while i < len(lines) and '|' in lines[i]:
                row_data = [c.strip() for c in lines[i].strip('|').split('|')]
                rows.append(row_data)
                i += 1
            
            table = doc.add_table(rows=len(rows)+1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Header
            hdr_cells = table.rows[0].cells
            for idx, h in enumerate(headers):
                if idx < len(hdr_cells):
                    hdr_cells[idx].text = h
                    pad = hdr_cells[idx].paragraphs[0]
                    pad.runs[0].font.bold = True
            
            # Rows
            for r_idx, row in enumerate(rows):
                row_cells = table.rows[r_idx+1].cells
                for c_idx, cell_text in enumerate(row):
                    if c_idx < len(row_cells):
                        _add_formatted_text(row_cells[c_idx].paragraphs[0], cell_text)
            
            doc.add_paragraph() # Spacer
            continue # Already advanced i

        # Regular Text
        elif line:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

        i += 1

    doc.save(OUTPUT_FILE)
    print(f"Successfully generated DOCX: {OUTPUT_FILE}")

def _add_formatted_text(paragraph, text):
    """Simple parser for **bold** text"""
    parts = text.split('**')
    for idx, part in enumerate(parts):
        run = paragraph.add_run(part)
        if idx % 2 == 1: # Odd parts are bold
            run.bold = True

if __name__ == "__main__":
    create_report()
