"""
Convert feature_analysis.md to DOCX format with flowchart
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Read markdown file
with open(r'c:\Users\91896\.gemini\antigravity\brain\250ee9d8-6b4b-4f78-8606-60f4671feeea\feature_analysis.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()

# Add title
title = doc.add_heading('Vasundhara - Complete Feature & Algorithm Analysis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Process markdown content
lines = content.split('\n')
i = 0

while i < len(lines):
    line = lines[i].strip()
    
    # Skip the first title since we added it manually
    if i == 0 and line.startswith('# '):
        i += 1
        continue
    
    # Headers
    if line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    
    # Lists
    elif line.startswith('- [ ] ') or line.startswith('- '):
        text = line[6:] if line.startswith('- [ ] ') else line[2:]
        p = doc.add_paragraph(text, style='List Bullet')
    elif line.startswith('✅ '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
    
    # Code blocks
    elif line.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1
        if code_lines:
            p = doc.add_paragraph('\n'.join(code_lines))
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
    
    # Tables
    elif '|' in line and i + 1 < len(lines) and '|' in lines[i+1]:
        # Parse table
        headers = [cell.strip() for cell in line.split('|')[1:-1]]
        i += 2  # Skip separator line
        rows = []
        while i < len(lines) and '|' in lines[i]:
            row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
            rows.append(row)
            i += 1
        
        # Create table
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Add rows
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                table.rows[row_idx + 1].cells[col_idx].text = cell_data
        
        doc.add_paragraph()  # Add spacing after table
        i -= 1  # Adjust for outer loop increment
    
    # Regular paragraphs
    elif line:
        # Handle bold (**text**)
        if '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:  # Odd indices are bold
                    run.bold = True
        else:
            doc.add_paragraph(line)
    
    i += 1

# Add page break before flowchart
doc.add_page_break()

# Add System Architecture Flowchart section
doc.add_heading('System Architecture & Workflow Flowchart', 1)
doc.add_paragraph('The following diagram shows the complete system architecture, including both implemented and planned features:')
doc.add_paragraph()

# Find and add the flowchart image
flowchart_image = r'C:/Users/91896/.gemini/antigravity/brain/250ee9d8-6b4b-4f78-8606-60f4671feeea/vasundhara_system_flowchart_1764863196237.png'
try:
    doc.add_picture(flowchart_image, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    doc.add_paragraph(f'[Flowchart image could not be loaded: {str(e)}]')

doc.add_paragraph()

# Add legend
doc.add_heading('Legend:', 2)
legend_items = [
    ('✅ Implemented', 'Fully working features with complete functionality'),
    ('🚧 Partially Done', 'Structure exists, needs business logic implementation'),
    ('🔄 Simulated', 'Mock/rule-based implementation, needs real ML models'),
    ('❌ Not Started', 'Planned features not yet implemented')
]

for symbol, description in legend_items:
    p = doc.add_paragraph()
    run1 = p.add_run(symbol)
    run1.bold = True
    run2 = p.add_run(f': {description}')

# Add flowchart description
doc.add_paragraph()
doc.add_heading('Architecture Overview:', 2)
overview_text = """
The Vasundhara system follows a modern microservices architecture with three main layers:

1. **User Layer**: Three distinct user types (Household, Shopkeeper, Admin) with role-based access control.

2. **Frontend Layer**: Next.js application serving different interfaces based on user role, including the comprehensive admin panel with 9 dedicated management pages.

3. **Backend Services**:
   - Express.js API handling business logic
   - FastAPI ML service for AI predictions
   - MongoDB for data persistence
   - Redis for caching and session management
   - RabbitMQ for asynchronous job processing

4. **External Integrations**: Email, SMS, push notifications, maps, and cloud storage services.

The diagram shows the complete data flow from user interactions through authentication, frontend components, API endpoints, ML services, and database operations.
"""
doc.add_paragraph(overview_text.strip())

# Save document
output_path = r'c:\Users\91896\Desktop\Vasundhara- Sugam Seva\Vasundhara_Complete_Analysis_With_Flowchart.docx'
doc.save(output_path)
print(f"DOCX file created successfully with flowchart at: {output_path}")
