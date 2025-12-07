"""
Convert project_languages.md to DOCX format
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Paths
INPUT_FILE = r'C:\Users\91896\.gemini\antigravity\brain\a0da7af2-194e-4dba-9882-8564e1741a2f\project_languages.md'
OUTPUT_FILE = r'c:\Users\91896\Desktop\Vasundhara- Sugam Seva\Vasundhara_Project_Languages.docx'

def convert_to_docx():
    print(f"Reading from: {INPUT_FILE}")
    
    try:
        with open(INPUT_FILE, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"Error: Input file not found at {INPUT_FILE}")
        return

    # Create document
    doc = Document()

    # Add main title
    title = doc.add_heading('Vasundhara Project - Language Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process markdown content
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        
        # Skip the original markdown title if present (since we added a docx title)
        if i == 0 and line.startswith('# '):
            i += 1
            continue

        # Headers
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)

        # Lists (Bullets)
        elif line.startswith('* ') or line.startswith('- '):
            text = line[2:] 
            # Handle bold inside bullets (simple handling)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)

        # Regular paragraphs
        elif line:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
        
        i += 1

    # Save document
    doc.save(OUTPUT_FILE)
    print(f"DOCX file created successfully at: {OUTPUT_FILE}")

def _add_formatted_text(paragraph, text):
    """Helper to handle basic **bold** formatting"""
    if '**' in text:
        parts = text.split('**')
        for idx, part in enumerate(parts):
            run = paragraph.add_run(part)
            if idx % 2 == 1:  # Odd indices are bold
                run.bold = True
    else:
        paragraph.add_run(text)

if __name__ == "__main__":
    convert_to_docx()
